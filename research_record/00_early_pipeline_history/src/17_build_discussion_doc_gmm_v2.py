from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\SDS2025_workspace\CASA_FYP\FYP")
IMG = ROOT / "outputs"
OUT = ROOT / "night_transport_RQ1_discussion_v2.docx"

PURPLE = "500778"
ACCENT = "8E6BB8"
BODY = "222222"
GREY = "666666"
LIGHT_PURPLE = "F4F0F8"
DISCUSS_FILL = "FBF3DE"
DISCUSS_BORDER = "C89B3C"
RISK_RED = "9A3D3D"


FIGS = {
    "rail_weekday_kdiag": IMG / "lu_rail_2000_featv2_FS" / "weekday_kdiag_v2.png",
    "rail_weekday_k3_profiles": IMG
    / "lu_rail_2000_featv2_FS"
    / "candidates"
    / "weekday_k3_profiles.png",
    "rail_weekday_k5_profiles": IMG
    / "lu_rail_2000_featv2_FS"
    / "candidates"
    / "weekday_k5_profiles.png",
    "rail_weekday_k4_map": IMG
    / "lu_rail_2000_featv2_FS"
    / "candidates"
    / "weekday_k4_map.png",
    "bus_weekday_k3_map": IMG
    / "busto_lsoa_2000_featv2_FS"
    / "candidates"
    / "weekday_k3_map.png",
    "bus_weekend_k3_map": IMG
    / "busto_lsoa_2000_featv2_FS"
    / "candidates"
    / "weekend_k3_map.png",
    "bus_weekday_kdiag": IMG / "busto_lsoa_2000_featv2_FS" / "weekday_kdiag_v2.png",
    "bus_weekend_kdiag": IMG / "busto_lsoa_2000_featv2_FS" / "weekend_kdiag_v2.png",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="CCCCCC", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_in):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=10.5, bold=False, italic=False, color=BODY):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_p(doc, parts=None, text=None, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    if text is not None:
        r = p.add_run(text)
        set_run_font(r)
    if parts:
        for value, opts in parts:
            r = p.add_run(value)
            set_run_font(
                r,
                size=opts.get("size", 10.5),
                bold=opts.get("bold", False),
                italic=opts.get("italic", False),
                color=opts.get("color", BODY),
            )
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=15 if level == 1 else 12, bold=True, color=PURPLE)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=8.5, italic=True, color=GREY)
    return p


def add_discuss(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.5)
    set_cell_shading(cell, DISCUSS_FILL)
    set_cell_border(cell, DISCUSS_BORDER, "8")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(3)
    r = p1.add_run("For discussion")
    set_run_font(r, size=9.5, bold=True, color="7A5A12")
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.12
    r2 = p2.add_run(text)
    set_run_font(r2, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_note(doc, label, text, fill=LIGHT_PURPLE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.5)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "D7CBE5", "4")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + ": ")
    set_run_font(r, size=9.5, bold=True, color=PURPLE)
    r2 = p.add_run(text)
    set_run_font(r2, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_img(doc, path, width, caption):
    if not path.exists():
        add_note(doc, "Missing figure", str(path), fill="FCE8E6")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_two_up_images(doc, left_path, left_caption, right_path, right_caption):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [3.2, 3.2]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_border(cell, "FFFFFF", "0")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for i, path in enumerate([left_path, right_path]):
        cell = table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path.exists():
            p.add_run().add_picture(str(path), width=Inches(3.05))
        else:
            r = p.add_run("Missing figure")
            set_run_font(r, size=9, color=RISK_RED, bold=True)
    for i, caption in enumerate([left_caption, right_caption]):
        cell = table.cell(1, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        set_run_font(r, size=8, italic=True, color=GREY)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, PURPLE)
        set_cell_border(cell, "FFFFFF", "4")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=8.5, bold=True, color="FFFFFF")
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_width(cell, widths[i])
            set_cell_border(cell)
            if ri % 2 == 0:
                set_cell_shading(cell, LIGHT_PURPLE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_run_font(r, size=8.5, bold=(i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run_font(r)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BODY)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("RQ1 discussion brief v2 (GMM) | revised for 2026-06-24 meeting")
    set_run_font(r, size=8, color=GREY)


def build():
    doc = Document()
    configure_doc(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Night-time Public Transport & Spatial Equity in London")
    set_run_font(r, size=17, bold=True, color=PURPLE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("RQ1 clustering discussion brief v2: evidence, uncertainty, and meeting decisions")
    set_run_font(r, size=12, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Fangzheng | UCL CASA x TfL | revised for supervisor discussion, 2026-06-24")
    set_run_font(r, size=9.5, italic=True, color=GREY)

    add_note(
        doc,
        "Purpose",
        "This is a pre-meeting discussion brief, not a final results report. "
        "The aim is to show what the current RQ1 pipeline has produced, where the evidence is strong, "
        "and where methodological decisions are still needed before cluster labels are finalised.",
    )

    add_heading(doc, "0. What changed in this revised brief")
    add_bullet(
        doc,
        "The window argument is reframed as a sensitivity trade-off, not as a settled claim that "
        "the evening peak simply 'contaminated' the clustering."
    )
    add_bullet(
        doc,
        "The central interpretation is moved toward a continuum-plus-special-cases framing: most "
        "locations follow similar night-decline profiles, while small but meaningful types appear only "
        "under higher K or targeted detection."
    )
    add_bullet(
        doc,
        "Figures are now used as evidence for specific meeting questions: K ambiguity, typology mismatch, "
        "the profile-plot limitation, bus aggregation, and low/no-service LSOAs."
    )
    add_bullet(
        doc,
        "Claims that belong to RQ2, especially equity interpretation, are kept as 'equity-relevant' rather "
        "than treated as completed findings."
    )

    add_heading(doc, "1. Evidence map for the meeting")
    add_p(
        doc,
        text=(
            "The table below shows where current analysis figures should be used. The goal is not to "
            "decorate the document, but to make each figure answer a specific methodological question."
        ),
    )
    add_simple_table(
        doc,
        ["Issue", "Figure / evidence to show", "What it supports", "Meeting ask"],
        [
            [
                "Window",
                "No current figure proves this directly; needs 18:00 vs 20:00 sensitivity panel",
                "20:00 removes some commute signal, but may also remove valid early-night activity",
                "Confirm primary window and whether sensitivity is required",
            ],
            [
                "K choice",
                "Rail weekday FS K-diagnostic",
                "Silhouette, DB, CH, and stability do not agree",
                "Choose K rule or abandon single global K",
            ],
            [
                "Typology mismatch",
                "Rail weekday K=3 vs K=5 profiles",
                "Low K gives broad main families; higher K reveals a small nightlife-like profile",
                "Decide between global clusters and two-tier typology",
            ],
            [
                "Feature/plot mismatch",
                "Profile plots plus feature signatures",
                "Direction balance (A_net) is used by the model but hidden in direction-normalised profiles",
                "Add total-normalised/asymmetry profile plots",
            ],
            [
                "Bus aggregation",
                "Bus FS LSOA maps and K-diagnostics",
                "Bus clusters are smoother and less sharply separated; FS maps also show a combined no/low-data layer",
                "Confirm bus unit and treatment of low/no-service areas",
            ],
        ],
        [1.1, 1.8, 2.05, 1.55],
    )

    add_heading(doc, "2. Window choice: not simply 'evening peak contamination'")
    add_p(
        doc,
        parts=[
            (
                "Current working choice: clustering starts at 20:00. ",
                {"bold": True},
            ),
            (
                "The earlier rationale was that 18:00-20:00 contains a strong evening-peak / commute-tail "
                "component. That remains plausible, but it should be treated as a trade-off rather than a "
                "settled correction. Including 18:00-20:00 may produce clearer separation, but that clarity "
                "may partly reflect conventional commuting rather than night-specific behaviour. Starting at "
                "20:00 reduces commute-driven separation, but it may also remove real early-evening night-work "
                "and leisure signals.",
                {},
            ),
        ],
    )
    add_note(
        doc,
        "Important reframing",
        "If the 20:00-start solution is less clean, this does not necessarily mean the model has failed. "
        "It may mean that after the evening peak is removed, night-time demand profiles form a weaker continuum "
        "rather than a few sharply separated types.",
        fill="FCE8E6",
    )
    add_discuss(
        doc,
        "Should 20:00 remain the primary RQ1 window, with 18:00-start results reported as a sensitivity check? "
        "Or should the thesis keep the LNWC-aligned 18:00-06:00 frame and explicitly acknowledge that one "
        "cluster dimension is early-evening commuting?"
    )

    add_heading(doc, "3. Current pipeline and what is already achieved")
    add_p(
        doc,
        text=(
            "The current RQ1 pipeline has moved beyond prototype data. Rail uses the five NUMBAT day-types "
            "and retains 270 LU-including stations. The latest outputs shown here are the FS variants: "
            "rail weekday = MON+TWT+SUN and weekend = FRI+SAT; bus weekday = Weekday+Sunday and weekend = "
            "Saturday-only because Friday is not separable in BUSTO. Bus uses the full set of 12 BUSTO files "
            "and is currently aggregated to LSOA for the clustering outputs shown here. The current feature set is behavioural "
            "rather than high-dimensional raw shares: A_net, F_flip, t_median, persist_00, hump_22, and dawn "
            "where relevant. GMM with diagonal covariance is the working model; KMeans and Ward are retained "
            "as checks, not as the main output."
        ),
    )
    add_discuss(
        doc,
        "The meeting does not need to decide whether the pipeline works. It needs to decide how the current "
        "outputs should be framed methodologically before labels are finalised."
    )

    add_heading(doc, "4. K is ambiguous under GMM")
    add_p(
        doc,
        text=(
            "The internal indices do not point to a single K. In the rail weekday FS variant, silhouette is "
            "high at K=2 and K=5, Davies-Bouldin favours K=4, and bootstrap stability is high for several K. "
            "This pattern is consistent with a broad continuum plus small pockets, rather than a clean set of "
            "globally separable groups."
        ),
    )
    add_img(
        doc,
        FIGS["rail_weekday_kdiag"],
        6.25,
        "Figure 1. Rail weekday FS K-diagnostics. The indices disagree, so K selection is not a mechanical step.",
    )
    add_discuss(
        doc,
        "Should K be selected by a single index, by interpretability, or should RQ1 avoid a single global K "
        "and instead report main profile families plus separately identified special cases?"
    )

    add_heading(doc, "5. The core issue: clusters are not the same as night-use types")
    add_p(
        doc,
        parts=[
            ("Main interpretation: ", {"bold": True}),
            (
                "the difficulty is not only the time window or the algorithm. Partitional clustering spends "
                "its K budget splitting the large homogeneous mass of smooth-decline stations into near-duplicate "
                "groups, while small but meaningful types only appear at higher K or through targeted detection.",
                {},
            ),
        ],
    )
    add_img(
        doc,
        FIGS["rail_weekday_k3_profiles"],
        5.7,
        "Figure 2. Rail weekday FS K=3 profiles. The main clusters are broad and partly similar.",
    )
    add_img(
        doc,
        FIGS["rail_weekday_k5_profiles"],
        5.25,
        "Figure 3. Rail weekday FS K=5 profiles. A small distinctive profile emerges, but most clusters remain close variants of smooth decline.",
    )
    add_note(
        doc,
        "Proposed meeting framing",
        "RQ1 may be stronger if framed as 'main profile families plus special cases' rather than as a single "
        "flat cluster typology. This is analytically more honest and preserves small substantively important "
        "types such as nightlife-like or airport-like patterns.",
    )
    add_discuss(
        doc,
        "Which strategy is acceptable for the dissertation: (a) low-K main types plus separate special-case "
        "detection, (b) peel off known special cases first, then cluster the remainder, or (c) describe a "
        "continuum with a special-case catalogue?"
    )

    add_heading(doc, "6. Feature engineering: useful, but needs a baseline check")
    add_p(
        doc,
        text=(
            "The engineered features improved separation compared with the raw temporal-share GMM, but they "
            "also inject behavioural assumptions. The most important new axis is A_net, because the raw "
            "direction-normalised profiles retained timing differences but lost the entry-vs-exit volume balance. "
            "This is defensible, but it should be checked against a more hands-off baseline such as raw-share + PCA "
            "or a compositional transformation."
        ),
    )
    add_note(
        doc,
        "Visualization gap",
        "The current profile plots normalise entry and exit separately, so they hide A_net. A cluster can be "
        "separated by the model but still look similar in the profile figure. Add a total-normalised profile or "
        "entry-minus-exit asymmetry plot before final reporting.",
        fill="FCE8E6",
    )
    add_discuss(
        doc,
        "Are the engineered features methodologically acceptable as the primary RQ1 representation, provided "
        "that a raw-share/PCA baseline is reported as a robustness check?"
    )

    add_heading(doc, "7. Algorithm and covariance choice")
    add_p(
        doc,
        text=(
            "GMM remains useful because it gives posterior probabilities and a probabilistic typology. However, "
            "the current evidence does not make the algorithm choice neutral: KMeans and Ward sometimes give "
            "higher silhouette, while GMM produces more imbalanced clusters and a less clear K. Full covariance "
            "is now numerically feasible in the reduced feature space, but diagonal covariance is retained as a "
            "conservative working choice because full covariance is not empirically supported in the current "
            "small-n setting."
        ),
    )
    add_discuss(
        doc,
        "Should GMM-diagonal remain the primary model, with KMeans/Ward as sensitivity checks, or should the "
        "simpler geometry of KMeans/Ward be preferred given the low-dimensional engineered feature space?"
    )

    add_heading(doc, "8. Weekday/weekend definition and the rail-bus mismatch")
    add_p(
        doc,
        text=(
            "The Fri+Sat weekend variant is conceptually strong for rail because Sunday has no Night Tube and "
            "appears closer to non-Night-Tube weekday behaviour. But BUSTO cannot isolate Friday: its day-types "
            "are Weekday, Saturday, and Sunday. Therefore the closest bus analogue makes weekend effectively "
            "Saturday-only. This is a data-forced cross-mode mismatch, not just a coding detail."
        ),
    )
    add_discuss(
        doc,
        "Is it acceptable to use the best behaviourally meaningful definition within each mode, even if rail "
        "and bus are not temporally identical? Or should RQ1 use a common but weaker weekday/weekend definition "
        "for cross-mode comparability?"
    )

    add_heading(doc, "9. Bus: aggregation, continuum, and low/no-service areas")
    add_p(
        doc,
        text=(
            "Bus clusters look smoother than rail clusters. This should not be read as stronger clustering. "
            "It may reflect LSOA aggregation: distinctive stop-level signals are averaged into area-level "
            "profiles, producing a more continuous spatial surface. Separately, low-service and no-service "
            "LSOAs should be retained as an equity-relevant layer, but not yet described as a final equity "
            "finding until they are compared with LNWC and socio-economic indicators in RQ2. The current FS "
            "candidate maps show a combined no/low-data layer; if the meeting needs the previous split between "
            "low service and no night service, that map should be regenerated under the FS definition."
        ),
    )
    add_img(
        doc,
        FIGS["bus_weekday_k3_map"],
        5.75,
        "Figure 4. Bus weekday FS K=3 candidate map, with unclustered no/low-data LSOAs shown in light grey.",
    )
    add_img(
        doc,
        FIGS["bus_weekday_kdiag"],
        5.8,
        "Figure 5. Bus weekday FS K-diagnostics. Bus separation is weaker and should be interpreted cautiously.",
    )
    add_discuss(
        doc,
        "Should bus RQ1 stay at LSOA level for direct linkage to LNWC, or should stop-level clustering be "
        "reintroduced for RQ1 and LSOA aggregation reserved for RQ2?"
    )

    add_heading(doc, "10. Current output examples (provisional, not labels)")
    add_p(
        doc,
        text=(
            "The maps below are useful to show that the pipeline now produces spatially interpretable outputs, "
            "but they should not be presented as final cluster labels. Their role in the meeting is to make the "
            "methodological trade-offs concrete."
        ),
    )
    add_two_up_images(
        doc,
        FIGS["rail_weekday_k4_map"],
        "Figure 6a. Rail weekday FS K=4 candidate map.",
        FIGS["bus_weekend_k3_map"],
        "Figure 6b. Bus weekend FS K=3 candidate map (Saturday-only weekend).",
    )

    add_heading(doc, "11. Recommended meeting decisions")
    for item in [
        "Confirm the primary temporal window and whether 18:00-start results should be reported as sensitivity analysis.",
        "Decide whether RQ1 should use a single global clustering solution or a two-tier typology.",
        "Confirm whether engineered behavioural features are acceptable as the main representation, with raw-share/PCA as robustness check.",
        "Confirm the primary model family: GMM-diagonal as main, or KMeans/Ward as main with GMM as sensitivity.",
        "Decide how to handle rail Fri+Sat weekend vs bus Saturday-only weekend.",
        "Decide whether bus clustering should remain LSOA-level for RQ1 or move back to stop-level with LSOA used in RQ2.",
        "Agree how absolute night activity should enter later: as a volume tier, a separate interpretive layer, or a sensitivity variant.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "12. Suggested figures still worth generating after the meeting")
    add_simple_table(
        doc,
        ["Figure", "Purpose", "Why it matters"],
        [
            [
                "18:00 vs 20:00 sensitivity panel",
                "Compare cluster diagnostics and profiles under both start times",
                "Tests whether the evening-peak argument is empirical rather than assumed",
            ],
            [
                "Total-normalised entry/exit profiles",
                "Make A_net visible in the same plot as timing",
                "Fixes the current mismatch between model features and profile figures",
            ],
            [
                "Entry-minus-exit asymmetry curves",
                "Show when each cluster is origin- or destination-dominant",
                "Better supports labels such as destination, residential-return, or airport-like",
            ],
            [
                "Cluster x volume-tier table/map",
                "Add total night activity without letting volume dominate clustering",
                "Bridges RQ1 profile typology to RQ2/RQ3 equity interpretation",
            ],
        ],
        [1.9, 2.35, 2.25],
    )

    doc.save(OUT)
    print(f"Saved {OUT}")


if __name__ == "__main__":
    build()
