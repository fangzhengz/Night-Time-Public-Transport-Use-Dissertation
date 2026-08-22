# -*- coding: utf-8 -*-
"""Render the RQ1 data/methods/validation brief to a styled .docx.

Purple/Calibri theme matching the existing discussion deliverables. Embeds the
back-filled K-diagnostic charts, per-K profiles and cluster maps produced by
script 18. Output: night_transport_RQ1_data_methods.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"D:\SDS2025_workspace\CASA_FYP\FYP")
OUT = ROOT / "outputs"
DEST = ROOT / "night_transport_RQ1_data_methods.docx"
PURPLE = RGBColor(0x50, 0x07, 0x78)
INK = RGBColor(0x2B, 0x2B, 0x2B)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = INK


def H(text, size=15, space_before=10):
    p = doc.add_paragraph(); p.space_before = Pt(space_before)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = PURPLE
    r.font.name = "Calibri"
    return p


def para(segs, space=4):
    """segs: list of (text, bold) or a plain string."""
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    p.paragraph_format.line_spacing = 1.12
    if isinstance(segs, str):
        segs = [(segs, False)]
    for text, bold in segs:
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(10.5); r.font.color.rgb = INK
    return p


def bullet(segs):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    if isinstance(segs, str):
        segs = [(segs, False)]
    for text, bold in segs:
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(10.5)


def fig(path, width=6.3, caption=None):
    if not Path(path).exists():
        para([("[缺图: %s]" % Path(path).name, True)]); return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = PURPLE


# ---- title ----
t = doc.add_paragraph(); t.space_after = Pt(2)
r = t.add_run("RQ1 — 数据、方法与检验过程说明"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = PURPLE
s = doc.add_paragraph()
rr = s.add_run("Night-time Public Transport & Spatial Equity in London · 供 2026-06-24 supervisor 会议讨论")
rr.italic = True; rr.font.size = Pt(9.5); rr.font.color.rgb = INK
para([("目的:现有讨论文档聚焦开放问题,本文补齐前提层——用了什么数据、做了什么处理、什么模型、做过哪些检验。"
       "数据范围与处理已定;聚类的算法/K/窗口/特征版本仍在讨论(见末节)。", False)])

# ---- 1 data ----
H("1. 数据来源与覆盖范围(范围是受限的)")
para([("Rail = NUMBAT(NBT24 系列,2024 秋季典型周)。", True),
      ("NUMBAT 覆盖 LU/LO/DLR/Elizabeth Line/Trams 全 TfL 轨道,但本研究只取伦敦地铁(LU):"
       "从 471 个 NUMBAT 车站中保留 mode set 含 LU 的站,得 270 个站为分析全集,"
       "主动排除 LO/DLR/EZL-only 与 39 个无闸机数据的电车站。每站 5 个日类型(MON/TWT/FRI/SAT/SUN)× entry+exit × 15-min。", False)])
bullet([("含义:rail 结论只对地铁成立,不能外推到 Overground/DLR/Elizabeth Line 的夜间使用。", False)])
bullet([("理由:与既有文献(Reades/Verma/Kimani/BtC)可比;单一模式票价、运营、闸机口径一致;地铁夜间(Night Tube)信号最相关。", False)])
para([("Bus = BUSTO(全网,4 个线路组 × 3 个日类型 = 12 个文件)。", True),
      ("原始粒度 route×direction×stop×15-min。本研究先对每个 STOPCODE 跨 route/direction 求和,"
       "再空间 join 到 LSOA、聚合到 LSOA 级别作为分析单元。日类型仅 Weekday/Saturday/Sunday(BUSTO 无独立 Friday)。", False)])
bullet([("含义:bus 分析单元是 LSOA(area-level),rail 是站点(point-level)——两侧单元不同,系有意为之(bus stop 太密)。", False)])
bullet([("代价:LSOA 聚合带来 MAUP,把站点层异质性抹平成平滑连续体——这是 bus 类型偏少、轮廓偏低的结构性原因,非算法问题。", False)])
para([("模态不对称是 finding,不是 bug。", True),
      ("rail=点/bus=面、粒度 15-min vs hourly、日类型 5 vs 3——这些差异保留,RQ2 通过对 LNWC/Census 交叉制表调和,而不合并成单一特征向量。", False)])

# ---- 2 preprocessing ----
H("2. 预处理与特征构造")
bullet([("窗口:", True), ("主夜间窗口 20:00 起(当前主线);rail weekday 20:00–01:00、weekend 延到 06:00(跨午夜续接);bus 20:00–06:00。"
        "预处理保留 18:00–06:00,聚类才从 20:00 切入以剔除 PM 通勤尾巴——代价是降低分离度(通勤本是强判别信号),属已知权衡。", False)])
bullet([("粒度:", True), ("rail = 15-min;bus = hourly(BIN=60)。", False)])
bullet([("低活动过滤:", True), ("单元夜间总活动 < 50 即剔除(MIN_TOTAL)。", False)])
bullet([("归一化:", True), ("每单元×每日类型×每方向对窗口内总量归一化为时间占比(聚类基于时序形状而非规模)。"
        "raw-share 版用 per-vector 归一化(每个'日类型×方向'子块各自和为 1)——与 BtC paper(Eq.1–2)一致,不是 joint。", False)])
para([("特征版本(三代):", True)])
bullet([("v1 — raw-share:", True), ("直接拿归一化曲线聚类(rail ~80–112 维 / bus 20–24 维),复刻 BtC/Kimani。问题:共享'夜间衰减包络'主导距离、簇近重复、silhouette 低。", False)])
bullet([("v2 — 行为特征工程:", True), ("压成 ~6 个可解释标量(A_net 净方向性、F_flip 方向翻转、t_median、persist_00、hump_22、dawn),z-score 后聚类。把 v1 丢失的净方向性显式建模。", False)])
bullet([("FS — 日定义变体:", True), ("weekend=Fri+Sat、weekday=Mon+TWT+Sun(Sunday 无 Night Tube)。bus 无法镜像(无 Friday)。", False)])

# ---- 3 model ----
H("3. 聚类模型与设定")
para([("主模型:GMM(高斯混合)。", True),
      ("当前用对角协方差(covariance_type='diag'),n_init≥15–20、reg_covar=1e-6、random_state=42、EM tol=1e-6;"
       "每单元输出 max_posterior 与 entropy 作为分配置信度。", False)])
para([("与文献的偏离(需在写作中交代):", True),
      ("Kimani 用 full 协方差、BtC 用 BIC 搜索 14 种协方差(含 full VVV);本项目在高维 raw-share 下实测 full 的 BIC 远差于 diag"
       "(rail wd K=4:diag 414 vs full 2024,p≫n 参数爆炸),故保留 diag。", False)])
para([("K 选择:K=2..15 全扫,三指标 + BIC,不自动选 K。", True),
      ("Silhouette(高好)、Calinski-Harabasz(高好)、Davies-Bouldin(低好);GMM 另记 BIC/AIC、log-likelihood、簇大小、max_posterior、entropy。"
       "三指标一致 → 共识 K;不一致 → 报告全部候选,留待人工/会议决定(GMM 下 K 本就不稳)。", False)])
para([("交叉模型:KMeans、Ward 仅作 K-诊断中的 cross-model 校验,不作主结论。", False)])

# ---- 4 validation ----
H("4. 已完成的检验过程")
para([("围绕'簇是否真实、K 是否站得住、特征是否有效'已做:", False)])
bullet([("(1) K-阶梯诊断(K=2..15):", True), ("每个(模态×日类型×特征版本)产出 silhouette/CH/DB(GMM 加 BIC)随 K 的曲线与 CSV。", False)])
bullet([("(2) Cross-model 校验:", True), ("同一矩阵上比较 GMM/KMeans/Ward 的 silhouette——低维 v2 下 KMeans 高于 GMM,换算法会移动结果。", False)])
bullet([("(3) Bootstrap 稳定性(GMM ARI):", True), ("重采样计算 Adjusted Rand Index,衡量簇标签在扰动下的稳定性。", False)])
bullet([("(4) 后验置信度:", True), ("逐单元 max_posterior 与 entropy;高熵/低 max-posterior 提示分配不确定,也是探测特殊型的线索。", False)])
bullet([("(5) 特征空间可视化:", True), ("v2 的 A–F scatter;FINAL 版的 PCA 投影。", False)])
bullet([("(6) 无监督 sanity check:", True), ("Heathrow 3 站凭极端 F_flip+dawn 自动成簇,与 BtC paper Cluster 8(Heathrow, n=3)独立吻合。", False)])
bullet([("(7) 簇画像与代表单元:", True), ("每候选 K 输出时序 profile、cluster summary、representative units、(rail)entry−exit 不对称图。", False)])

H("4.1 对 FINAL 版补做的 K 诊断(2026-06-23 回填)", size=12)
para([("此前 ", False), ("lu_rail_FINAL_diag", True), (" 与 ", False), ("busto_lsoa_FINAL_gmm_hourly", True),
      (" 只给了单一硬选 K 的产物、缺 K 评估图,K 的来源无据。现已在其原始 raw-share 矩阵上重算 K=2..15 全套诊断"
       "(GMM/KMeans/Ward silhouette、CH、DB、BIC、bootstrap ARI),并对 4–8 个候选 K 重新产出标签/画像/空间图。"
       "重拟合对既有标签 ARI:rail weekday 0.82、bus weekday 0.92(高度复现,确认即原 pipeline);"
       "rail weekend 0.64、bus weekend 0.46(周末不稳)。", False)])
para([("关键发现:现有 FINAL 的 K(weekday=7 / weekend=8)不被任何内部指标支持。", True),
      ("Rail weekday——silhouette 三模型一致在 K=2 最高后单调下滑、CH 峰在 K=4、稳定性 ARI 峰在 K=5、BIC 无内部极值;"
       "Rail weekend——silhouette/CH 峰在 K=2、ARI 峰在 K=3。所有指标都指向远小于 7/8 的 K(2–5)。"
       "per-K 画像亦显示多数簇共享同一条夜间衰减包络,仅极少数小簇真正不同——即高维 raw-share 的共模主导。"
       "结论不是立即改 K,而是把这一不一致透明提交会议讨论。", False)])

fig(OUT/"lu_rail_FINAL_diag"/"weekday"/"weekday_kdiag.png", caption="图 1 · RAIL weekday FINAL(raw-share)K 诊断:silhouette/CH 偏好小 K,K=7 无支撑")
fig(OUT/"lu_rail_FINAL_diag"/"weekend"/"weekend_kdiag.png", caption="图 2 · RAIL weekend FINAL K 诊断:指标偏好 K=2–3,K=8 无支撑")
fig(OUT/"busto_lsoa_FINAL_gmm_hourly"/"weekday"/"weekday_kdiag.png", caption="图 3 · BUS weekday FINAL(LSOA hourly)K 诊断")
fig(OUT/"lu_rail_FINAL_diag"/"weekday"/"candidates"/"weekday_k5_profiles.png", width=4.6,
    caption="图 4 · RAIL weekday K=5 簇画像:多数簇共享衰减包络,仅 C3(n≈16)明显不同")
fig(OUT/"lu_rail_FINAL_diag"/"weekday"/"candidates"/"weekday_k5_map.png", width=5.6,
    caption="图 5 · RAIL weekday K=5 站点空间分布(按 NLC join 坐标,270/270 全匹配)")
fig(OUT/"busto_lsoa_FINAL_gmm_hourly"/"weekday"/"candidates"/"weekday_k4_map.png", width=5.6,
    caption="图 6 · BUS weekday K=4 LSOA 空间分布")

para([("尚未做(属 RQ2 或待定):", True),
      ("Random Forest 特征重要性、socio-demographic 验证(对 LNWC/Census/IMD)、catchment 定义(rail 1200m Voronoi-clipped / bus 待定)。", False)])

# ---- 5 lineage ----
H("5. 与参考文献方法的关系")
bullet([("v1 raw-share = BtC/Kimani 血统;", True), ("其中 lu_rail_FINAL_diag(15-min、entry+exit、per-vector、GMM)最贴 BtC,差异仅在夜窗(非全天)+ 对角协方差;"
        "busto_lsoa_FINAL_gmm_hourly 范式同向,但 hourly 粒度偏离 BtC 的 15-min。", False)])
bullet([("entry+exit 双向沿用 BtC", True), ("(Kimani 仅用 entry)。", False)])
bullet([("per-vector 归一化与 BtC 一致", True), ("(注意:不是 joint)。", False)])
bullet([("三指标选 K 沿用 Kimani;", True), ("BtC 实际用 BIC——两篇母本在 K 选择上不一致,写作时需分别归属。", False)])
bullet([("v2 工程特征是对两篇母本 raw-share 范式的共同突破", True), ("(针对夜窗低信噪比下的维度灾难/共模主导)。", False)])

# ---- 6 open ----
H("6. 仍待会议定夺")
for i, txt in enumerate([
    "特征版本:raw-share(贴文献)vs 工程特征 v2(判别力高但有先验注入风险)——建议补 raw-share→(CLR)→PCA→GMM 数据驱动中间路线对照。",
    "算法:GMM vs KMeans/Ward(结果会移动)。",
    "K:GMM 下三指标/稳定性指向不同 K,且 FINAL 的 7/8 与诊断冲突(见 4.1)。",
    "窗口:20:00 vs 18:00(18:00 更清晰但清晰来自通勤,非夜间)。",
    "周末定义 + rail-bus 时间口径错配(rail Fri-Sat vs bus Sat-only)。",
    "协方差:diag vs full(文献用 full,本项目 BIC 偏好 diag)。",
    "流量/体量未用:特征皆 scale-free;equity 关心体量,需作单独图层(cluster × volume tier)带回。",
], 1):
    bullet([("%d. " % i, True), (txt, False)])

doc.save(str(DEST))
print("saved:", DEST)
