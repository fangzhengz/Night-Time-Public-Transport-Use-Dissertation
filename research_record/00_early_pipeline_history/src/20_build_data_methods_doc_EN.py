# -*- coding: utf-8 -*-
"""English version of the RQ1 data/methods/validation brief (.docx).

Purple/Calibri theme matching the discussion deliverables. Embeds the
back-filled K-diagnostic charts, per-K profiles and cluster maps from script 18.
Night window stated as 18:00-06:00 (provisional; still under discussion).
Output: night_transport_RQ1_data_methods_EN.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"D:\SDS2025_workspace\CASA_FYP\FYP")
OUT = ROOT / "outputs"
DEST = ROOT / "night_transport_RQ1_data_methods_EN.docx"
PURPLE = RGBColor(0x50, 0x07, 0x78)
INK = RGBColor(0x2B, 0x2B, 0x2B)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = INK


def H(text, size=15, space_before=10):
    p = doc.add_paragraph(); p.space_before = Pt(space_before)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = PURPLE
    r.font.name = "Calibri"
    return p


def para(segs, space=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    p.paragraph_format.line_spacing = 1.12
    if isinstance(segs, str):
        segs = [(segs, False)]
    for text, bold in segs:
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(10.5); r.font.color.rgb = INK
    return p


def bullet(segs):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    if isinstance(segs, str):
        segs = [(segs, False)]
    for text, bold in segs:
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(10.5)


def fig(path, width=6.3, caption=None):
    if not Path(path).exists():
        para([("[missing figure: %s]" % Path(path).name, True)]); return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = PURPLE


# ---- title ----
t = doc.add_paragraph(); t.space_after = Pt(2)
r = t.add_run("RQ1 — Data, Methods and Validation"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = PURPLE
s = doc.add_paragraph()
rr = s.add_run("Night-time Public Transport & Spatial Equity in London · briefing for the 2026-06-24 supervisor meeting")
rr.italic = True; rr.font.size = Pt(9.5); rr.font.color.rgb = INK
para([("Purpose: the existing discussion document focuses on open problems; this note fills in the premises — "
       "what data are used, what processing was applied, which models, and what validation has been done. "
       "Data scope and processing are settled; the clustering algorithm / K / time window / feature version "
       "are still under discussion (see final section).", False)])

# ---- 1 data ----
H("1. Data sources and coverage (the scope is deliberately restricted)")
para([("Rail = NUMBAT (NBT24 series, typical autumn week 2024). ", True),
      ("NUMBAT covers the full TfL rail estate (LU / LO / DLR / Elizabeth Line / Trams), but this study uses "
       "London Underground (LU) only: of the 471 NUMBAT stations we keep those whose mode set includes LU, "
       "yielding 270 stations as the analysis universe. This deliberately excludes LO/DLR/EZL-only stations and "
       "the 39 tram stations that have no gateline data. Each station has 5 day types (MON/TWT/FRI/SAT/SUN) "
       "× entry+exit × 15-min bins.", False)])
bullet([("Implication: ", True), ("rail findings hold for the Underground only and cannot be extrapolated to "
        "Overground/DLR/Elizabeth Line night-time use.", False)])
bullet([("Rationale: ", True), ("comparability with prior work (Reades / Verma / Kimani / BtC); a single mode keeps "
        "fare structure, operations and gateline conventions consistent; the Underground carries the most relevant "
        "night-time (Night Tube) signal.", False)])
para([("Bus = BUSTO (full network: 4 route-group CSVs × 3 day types = 12 files). ", True),
      ("Raw grain is route × direction × stop × 15-min. We first sum each STOPCODE across routes/directions to a "
       "stop-level demand, then spatially join stops to LSOAs and aggregate to LSOA level as the analysis unit. "
       "Day types are Weekday / Saturday / Sunday only (BUSTO has no separate Friday).", False)])
bullet([("Implication: ", True), ("the bus analysis unit is the LSOA (area-level) whereas rail is the station "
        "(point-level) — the units differ by design (bus stops are too granular; interchanges hold many stops).", False)])
bullet([("Cost: ", True), ("LSOA aggregation introduces MAUP, smoothing stop-level heterogeneity into a continuum — "
        "this is the structural reason bus yields fewer types and lower silhouettes, not an algorithmic failure.", False)])
para([("Modal asymmetry is a finding, not a bug. ", True),
      ("rail = point / bus = area, granularity 15-min vs hourly, day types 5 vs 3 — these differences are preserved "
       "and reconciled at RQ2 via cross-tabulation against LNWC/Census, rather than merged into one feature vector.", False)])

# ---- 2 preprocessing ----
H("2. Pre-processing and feature construction")
bullet([("Time window (PROVISIONAL — still disputed): ", True),
        ("18:00–06:00 is used for now. Rail weekday is truncated to ~01:00 because trains stop running "
         "(post-closure bins are structurally zero, not imputed); rail weekend extends to 06:00 where Night Tube "
         "operates; bus runs 18:00–06:00 throughout. A 20:00 start was also trialled to drop the PM-commute tail "
         "(a generic, non-night signal), but that reduces cluster separation because the commute is itself a strong "
         "discriminator — hence the window remains an open question (see final section).", False)])
bullet([("Granularity: ", True), ("rail = 15-min; bus = hourly (BIN=60).", False)])
bullet([("Low-activity filter: ", True), ("a unit is dropped if its total night activity < 50 (MIN_TOTAL).", False)])
bullet([("Normalisation: ", True), ("per unit × day type × direction, counts are normalised to temporal shares "
        "(clustering targets profile shape, not volume). The raw-share version uses per-vector normalisation "
        "(each 'day-type × direction' block sums to 1) — consistent with BtC (Eq. 1–2), not joint normalisation.", False)])
para([("Feature versions (three generations):", True)])
bullet([("v1 — raw-share: ", True), ("cluster directly on the normalised per-bin curves (rail ~80–112 dims / bus 20–24 dims); "
        "reproduces BtC/Kimani. Problem: a shared 'night-decline envelope' dominates distance → near-duplicate "
        "clusters, low silhouette.", False)])
bullet([("v2 — behavioural feature engineering: ", True), ("collapse to ~6 interpretable scalars (A_net net directionality, "
        "F_flip direction flip, t_median, persist_00, hump_22, dawn), z-scored before clustering; explicitly models the "
        "net directionality that v1 discards.", False)])
bullet([("FS — day-definition variant: ", True), ("weekend = Fri+Sat, weekday = Mon+TWT+Sun (Sunday has no Night Tube). "
        "Bus cannot mirror this (no separate Friday).", False)])

# ---- 3 model ----
H("3. Clustering model and settings")
para([("Primary model: Gaussian Mixture Model (GMM). ", True),
      ("Currently diagonal covariance (covariance_type='diag'), n_init ≥ 15–20, reg_covar=1e-6, random_state=42, "
       "EM tol=1e-6; per-unit max_posterior and entropy are recorded as assignment-confidence diagnostics.", False)])
para([("Deviation from the references (to be stated in the write-up): ", True),
      ("Kimani uses full covariance; BtC searches 14 covariance parameterisations (incl. full VVV) via BIC. In our "
       "high-dimensional raw-share setting full covariance is heavily penalised by BIC relative to diagonal "
       "(rail weekday K=4: diag 414 vs full 2024, i.e. p ≫ n over-parameterisation), so diagonal is retained.", False)])
para([("K selection: full sweep K = 2..15, three indices + BIC, no auto-pick. ", True),
      ("Silhouette (higher better), Calinski-Harabasz (higher better), Davies-Bouldin (lower better); for GMM also "
       "BIC/AIC, log-likelihood, cluster sizes, max_posterior, entropy. If the three indices agree → consensus K; "
       "if they disagree → report all candidates and leave the choice to inspection / the meeting (K is intrinsically "
       "unstable under GMM here).", False)])
para([("Cross-model: KMeans and Ward are used only as cross-checks in the K-diagnostics, not for primary conclusions.", False)])

# ---- 4 validation ----
H("4. Validation performed to date")
para([("Around the questions 'are the clusters real, is K defensible, are the features effective':", False)])
bullet([("(1) K-ladder diagnostics (K=2..15): ", True), ("per (mode × day type × feature version), silhouette/CH/DB "
        "(plus BIC for GMM) as a function of K, with CSVs.", False)])
bullet([("(2) Cross-model check: ", True), ("GMM vs KMeans vs Ward silhouettes on the same matrix — under the low-dim v2 "
        "features KMeans exceeds GMM, i.e. switching algorithm moves the result.", False)])
bullet([("(3) Bootstrap stability (GMM ARI): ", True), ("resampling with Adjusted Rand Index to measure label stability "
        "under perturbation.", False)])
bullet([("(4) Posterior confidence: ", True), ("per-unit max_posterior and entropy; high entropy / low max-posterior flags "
        "uncertain assignments and helps surface special types.", False)])
bullet([("(5) Feature-space visualisation: ", True), ("A–F scatter for v2; PCA projection for the FINAL (raw-share) runs.", False)])
bullet([("(6) Unsupervised sanity check: ", True), ("Heathrow's 3 stations separate automatically on extreme F_flip+dawn, "
        "independently matching BtC's Cluster 8 (Heathrow, n=3).", False)])
bullet([("(7) Cluster profiles and representative units: ", True), ("per candidate K: temporal profiles, cluster summaries, "
        "representative units, and (rail) entry−exit asymmetry plots.", False)])

H("4.1 Back-filled K-diagnostics for the FINAL runs (added 2026-06-23)", size=12)
para([("Previously ", False), ("lu_rail_FINAL_diag", True), (" and ", False), ("busto_lsoa_FINAL_gmm_hourly", True),
      (" shipped deliverables for a single hard-picked K but lacked any K-evaluation chart, so the choice of K was "
       "unjustified. We have now recomputed the full K=2..15 diagnostics on their original raw-share matrices "
       "(GMM/KMeans/Ward silhouette, CH, DB, BIC, bootstrap ARI) and re-generated labels/profiles/maps for 4–8 "
       "candidate K. Refit-vs-existing-label agreement (ARI): rail weekday 0.82, bus weekday 0.92 (high reproduction, "
       "confirming this is the original pipeline); rail weekend 0.64, bus weekend 0.46 (weekend is unstable).", False)])
para([("Key finding: the existing FINAL K (weekday = 7 / weekend = 8) is supported by no internal index. ", True),
      ("Rail weekday — silhouette (all three models) peaks at K=2 then declines monotonically, CH peaks at K=4, "
       "stability ARI peaks at K=5, and BIC has no interior optimum; rail weekend — silhouette/CH peak at K=2, ARI at K=3. "
       "Every index points to a much smaller K (2–5). The per-K profiles likewise show that most clusters share one "
       "night-decline envelope, with only a few small clusters genuinely distinct — i.e. the common-mode domination of "
       "high-dimensional raw shares. The conclusion is not to change K now, but to surface this inconsistency "
       "transparently for the meeting.", False)])

fig(OUT/"lu_rail_FINAL_diag"/"weekday"/"weekday_kdiag.png", caption="Fig. 1 · RAIL weekday FINAL (raw-share) K-diagnostics: silhouette/CH favour small K; K=7 unsupported")
fig(OUT/"lu_rail_FINAL_diag"/"weekend"/"weekend_kdiag.png", caption="Fig. 2 · RAIL weekend FINAL K-diagnostics: indices favour K=2–3; K=8 unsupported")
fig(OUT/"busto_lsoa_FINAL_gmm_hourly"/"weekday"/"weekday_kdiag.png", caption="Fig. 3 · BUS weekday FINAL (LSOA hourly) K-diagnostics")
fig(OUT/"lu_rail_FINAL_diag"/"weekday"/"candidates"/"weekday_k5_profiles.png", width=4.6,
    caption="Fig. 4 · RAIL weekday K=5 cluster profiles: most clusters share the decline envelope; only C3 (n≈16) differs")
fig(OUT/"lu_rail_FINAL_diag"/"weekday"/"candidates"/"weekday_k5_map.png", width=5.6,
    caption="Fig. 5 · RAIL weekday K=5 station map (coords joined by NLC, 270/270 matched)")
fig(OUT/"busto_lsoa_FINAL_gmm_hourly"/"weekday"/"candidates"/"weekday_k4_map.png", width=5.6,
    caption="Fig. 6 · BUS weekday K=4 LSOA map")

para([("Not yet done (RQ2 or pending): ", True),
      ("Random Forest feature importance, socio-demographic validation (against LNWC/Census/IMD), and catchment "
       "definition (rail 1200m Voronoi-clipped / bus open).", False)])

# ---- 5 lineage ----
H("5. Relationship to the reference methods")
bullet([("v1 raw-share = BtC/Kimani lineage; ", True), ("of these, lu_rail_FINAL_diag (15-min, entry+exit, per-vector, GMM) "
        "is closest to BtC, differing only in the night window (vs full day) and diagonal covariance; "
        "busto_lsoa_FINAL_gmm_hourly is the same paradigm but its hourly granularity departs from BtC's 15-min.", False)])
bullet([("entry+exit follows BtC ", True), ("(Kimani clusters on entries only).", False)])
bullet([("per-vector normalisation matches BtC ", True), ("(note: not joint).", False)])
bullet([("three-index K selection follows Kimani; ", True), ("BtC actually uses BIC — the two source papers disagree on "
        "K selection, so attribute each choice explicitly in the write-up.", False)])
bullet([("v2 engineered features are a shared departure from both papers' raw-share paradigm ", True),
        ("(targeting the curse of dimensionality / common-mode domination under the low-SNR night window).", False)])

# ---- 6 open ----
H("6. Still to be decided at the meeting")
for i, txt in enumerate([
    "Feature version: raw-share (faithful to the literature) vs engineered v2 (more separable but risks injecting prior) — "
    "recommend adding a data-driven middle path, raw-share → (CLR) → PCA → GMM, for comparison.",
    "Algorithm: GMM vs KMeans/Ward (results shift on switching).",
    "K: the three indices / stability point to different K under GMM, and the FINAL 7/8 conflict with the diagnostics (see 4.1).",
    "Window: 18:00 (current, provisional) vs 20:00 — 18:00 is clearer but the clarity is commute-driven, not night-specific.",
    "Weekend definition + rail–bus temporal mismatch (rail Fri-Sat vs bus Saturday-only).",
    "Covariance: diag vs full (the literature uses full; here BIC favours diag).",
    "Flow/volume unused: all features are scale-free; equity cares about volume, so reintroduce it as a separate layer "
    "(cluster × volume tier).",
], 1):
    bullet([("%d. " % i, True), (txt, False)])

doc.save(str(DEST))
print("saved:", DEST)
