# -*- coding: utf-8 -*-
"""09 · Build the meeting results document (English, styled .docx).

Full narrative of the clean grouped analysis: data, preprocessing, filtering and
selection rationale, method choices, validation, and the cluster typologies with
tentative functional names. Reads figures + signature CSVs; writes the .docx to
the FYP root. Additive; does not modify the pipeline.
"""
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).resolve().parent))
import config as C

FIG = C.FIG
INTERP = C.OUT / "interpretation"
DEST = C.FYP / "night_transport_RQ1_results.docx"
PURPLE = RGBColor(0x50, 0x07, 0x78)
INK = RGBColor(0x2B, 0x2B, 0x2B)
GREY = RGBColor(0x66, 0x66, 0x66)

# tentative functional names per cluster (interpretation; code keeps cluster ids)
NAMES = {
    "rail_weekday": {0: "Central leisure/nightlife core (evening departure)",
                     1: "Residential-return mass (dominant)",
                     2: "Outer low-volume residential",
                     3: "Major interchange / terminus (balanced, high-volume)",
                     4: "Event/leisure destination — late (incl. O2)"},
    "rail_weekend": {0: "Inner-city destination, Night-Tube active",
                     1: "Residential weekend return, limited Night-Tube",
                     2: "Central nightlife origin core (West End) — signature type",
                     3: "Major termini (balanced, high-volume)"},
    "bus_weekday": {0: "High-volume town-centre / transport hub",
                    1: "Hospital / night-shift destination (overnight-active)",
                    2: "Outer suburban, evening-only",
                    3: "Central night core + airport (origin, strongest overnight)"},
    "bus_weekend": {0: "Airport / strong-overnight corridor (special)",
                    1: "Central (moderate overnight)",
                    2: "Outer suburban, minimal overnight",
                    3: "High-volume hubs / town centres"},
}
KPICK = {"rail_weekday": 5, "rail_weekend": 4, "bus_weekday": 4, "bus_weekend": 4}
COV = {"rail_weekday": "tied", "rail_weekend": "tied", "bus_weekday": "full", "bus_weekend": "full"}

doc = Document()
nm = doc.styles["Normal"]; nm.font.name = "Calibri"; nm.font.size = Pt(10.5); nm.font.color.rgb = INK


def H(text, size=15, before=12):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = PURPLE; r.font.name = "Calibri"


def para(segs, space=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space); p.paragraph_format.line_spacing = 1.12
    for text, bold in ([(segs, False)] if isinstance(segs, str) else segs):
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(10.5); r.font.color.rgb = INK


def bullet(segs):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    for text, bold in ([(segs, False)] if isinstance(segs, str) else segs):
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(10.5)


def fig(path, width=6.3, cap=None):
    path = Path(path)
    if not path.exists():
        para([("[missing figure: %s]" % path.name, True)]); return
    doc.add_picture(str(path), width=Inches(width)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(cap); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = PURPLE


def sig_table(ds):
    df = pd.read_csv(INTERP / f"{ds}_k{KPICK[ds]}_signature.csv")
    cols = [("cluster", "C"), ("n", "n"), ("net_dir", "net dir"), ("t_median", "t-median"),
            ("evening_18_20", "eve 18-20"), ("late_00plus", "overnight 00+"), ("hump_22_23", "hump 22-23")]
    t = doc.add_table(rows=1, cols=len(cols) + 1); t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for j, (_, lab) in enumerate(cols):
        hdr[j].paragraphs[0].add_run(lab).bold = True
    hdr[-1].paragraphs[0].add_run("tentative type").bold = True
    for _, r in df.iterrows():
        cells = t.add_row().cells
        for j, (k, _) in enumerate(cols):
            cells[j].paragraphs[0].add_run(str(r[k]))
        cells[-1].paragraphs[0].add_run(NAMES[ds].get(int(r["cluster"]), ""))
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for rn in p.runs:
                    rn.font.size = Pt(8.5)
    doc.add_paragraph()


# ===================== TITLE =====================
p = doc.add_paragraph(); r = p.add_run("RQ1 — Night-time Public Transport Use Profiles: Full Analysis")
r.bold = True; r.font.size = Pt(19); r.font.color.rgb = PURPLE
p = doc.add_paragraph(); r = p.add_run("Night-time Public Transport & Spatial Equity in London · clean re-analysis for the supervisor meeting (2026-06-24)")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY
para([("This document gives the complete, conclusion-oriented narrative of the rebuilt RQ1 clustering: data, "
       "preprocessing, the rationale for each filtering/selection choice, the analytical method, the validation "
       "performed, and the resulting station/area typologies. It supersedes the earlier problem-focused note; the "
       "pipeline was rebuilt from scratch in a single, reproducible flow (config-driven, one direction, isolated "
       "outputs).", False)])

# ===================== 1 DATA =====================
H("1. Data sources and coverage")
para([("Rail — TfL NUMBAT (NBT24 series, a typical autumn week, 2024). ", True),
      ("NUMBAT gives 15-minute station entry and exit counts for a typical Monday, Tuesday–Wednesday–Thursday "
       "(TWT), Friday, Saturday and Sunday. We restrict to the London Underground: of 471 NUMBAT stations we keep "
       "those whose mode set includes LU, giving 270 stations, and drop the 39 trams (no gateline data) and "
       "LO/DLR/Elizabeth-only stations.", False)])
bullet([("Why LU-only: ", True), ("comparability with the established London clustering literature (Reades, Verma, "
        "Kimani, Peiret-García et al.); a single mode keeps fare, operational and gateline conventions consistent; "
        "the Underground carries the clearest night-time (Night Tube) signal. Findings therefore apply to the "
        "Underground only.", False)])
para([("Bus — TfL BUSTO (full network, autumn typical demand). ", True),
      ("Delivered at route × direction × stop × 15-minute grain for Weekday, Saturday and Sunday (no separate "
       "Friday). For each stop we sum boardings and alightings across all routes and directions, then assign each "
       "stop to the LSOA that contains it (point-in-polygon, British National Grid) and aggregate to LSOA level.", False)])
bullet([("Spatial join result: ", True), ("of 21,515 bus stops with coordinates, 94.4% fall inside a London LSOA; "
        "1,138 BUSTO demand stops were unmatched (3.7% of rows, dropped). 4,111 of 4,994 LSOAs carry night bus "
        "activity.", False)])
bullet([("Why LSOA for bus: ", True), ("bus stops are too granular (~20k, many per interchange) to be a stable "
        "clustering unit, and LSOA aligns with the London Night Workers Classification for RQ2. The cost is MAUP — "
        "areal aggregation smooths stop-level heterogeneity into a continuum (the structural reason bus yields fewer, "
        "softer clusters than rail).", False)])
para([("Modal asymmetry is preserved as a finding: ", True), ("rail = station (point), bus = LSOA (area); rail "
       "15-minute, bus hourly; rail 5 day-types, bus 3. The two are clustered separately and reconciled later "
       "against LNWC/Census, not merged into one feature vector.", False)])

# ===================== 2 PREPROCESSING =====================
H("2. Window, granularity and filtering — with rationale")
bullet([("Night window 18:00–06:00 (evening peak retained). ", True), ("Dropping the 18:00–20:00 commute block was "
        "considered but rejected: it risks (i) appearing to engineer the clusters and (ii) excluding night workers "
        "who begin their shift commute in the early evening (supervisor guidance). A down-weighting sensitivity "
        "(§4) confirmed no benefit, so the block is kept at equal weight. Rail weekday closes ~01:00 (post-closure "
        "bins are structurally zero, not imputed); Fri/Sat run to 05:00 (Night Tube; NUMBAT's traffic day ends "
        "05:00); bus runs throughout.", False)])
bullet([("Granularity: ", True), ("rail 15-minute (validated, stable); bus hourly (15-min bins summed). Hourly is "
        "more stable for the areal bus signal and is consistent with TfL guidance; rail and bus therefore use "
        "different time slices by design.", False)])
bullet([("Low-volume filter: ", True), ("a unit is dropped if its total night activity < 50 (removes noise-dominated "
        "units).", False)])
bullet([("Normalisation = per-vector temporal shares (BtC). ", True), ("Each (direction) block is normalised to sum "
        "to 1, so clustering targets the SHAPE of the temporal profile, not volume. This is the standardisation; no "
        "StandardScaler is applied on top (it would distort the simplex). Absolute volume is retained separately for "
        "interpretation.", False)])

# ===================== 3 DAY GROUPING =====================
H("3. Day grouping — the central design choice")
para([("Clustering is run SEPARATELY per day-group, split by service regime:", False)])
bullet([("Rail: ", True), ("weekday = MON+TWT pooled (no Night Tube, ~01:00 close) | weekend = Friday+Saturday "
        "pooled (Night Tube, to 05:00). Sunday is excluded for now (no Night Tube; behaves like a weekday) — a "
        "documented scope limit.", False)])
bullet([("Bus: ", True), ("weekday = Weekday | weekend = Saturday+Sunday pooled (BUSTO has no separate Friday).", False)])
para([("Within a group the member days are POOLED (counts summed into one night profile, then per-vector "
       "normalised). Rationale — earlier we concatenated all day-types into one 288-dim vector and clustered once "
       "(BtC original); this DILUTED the Night-Tube overnight signal (the weekday/Sunday early-decline blocks "
       "dominated the distance, so the small overnight tail did not drive the clustering). Splitting by service "
       "regime lets each group's own signal lead. The result was a marked improvement in rail separation (weekend "
       "silhouette at K=2 rose to ~0.39) and, crucially, a clean nightlife cluster emerged (§5).", False)])

# ===================== 4 METHOD + VALIDATION =====================
H("4. Clustering method and validation")
para([("Model = Gaussian Mixture Model. ", True), ("Probabilistic soft assignment with per-unit max-posterior and "
       "entropy. Covariance type is selected by BIC across {spherical, diagonal, tied, full}: rail → tied, bus → "
       "full, a ranking that is stable across K. (Diagonal is permissible per supervisor but is not the BIC choice "
       "on these lower-dimensional grouped matrices.)", False)])
para([("Number of clusters K. ", True), ("BIC favours an over-parsimonious K=2 that masks small but substantively "
       "important types, so — with supervisor agreement that choosing a more interpretable K is standard practice — "
       "K is set on interpretability plus internal validation (silhouette, Davies-Bouldin, cluster sizes, bootstrap "
       "stability): rail weekday K=5, rail weekend K=4, bus K=4.", False)])
para([("Validation performed: ", True), ("(i) full K-diagnostics per dataset (silhouette / CH / DB / BIC / bootstrap "
       "ARI over K=2–12); (ii) BIC covariance selection; (iii) a rigorous evening down-weighting sensitivity — at "
       "each weight the whole BIC covariance×K selection was re-run, and measured on the TRUE (unweighted) "
       "overnight share, down-weighting did NOT improve overnight separation at any K, so the equal-weight window "
       "is retained; (iv) cluster signatures and representative units; (v) face-validity checks (the West-End "
       "nightlife cluster, Heathrow/Hatton Cross, and hospital stops all emerge where expected).", False)])
fig(FIG / "rail_weekend_kdiag.png", cap="Fig. 1 · Rail weekend K-diagnostics (BtC raw-share, GMM tied)")
fig(FIG / "rail_weekend_bic_grid.png", width=4.8, cap="Fig. 2 · BIC by covariance × K — tied selected for rail")

# ===================== 5 RESULTS =====================
H("5. Results — night-time use typologies")
para([("Reading the signatures. ", True), ("Net directionality net_dir is the dominant axis: in the night window "
       "net_dir > 0 (entry/boarding-dominant) = net OUTFLOW = trip ORIGIN (e.g. an activity district emptying after "
       "the evening); net_dir < 0 (exit/alighting-dominant) = net INFLOW = DESTINATION (e.g. residents arriving "
       "home). 'overnight 00+' is the share of the unit's night activity occurring after midnight (higher = more "
       "Night-Tube / night-owl activity). In the profile plots green = entry/boardings (outflow direction), red = "
       "exit/alightings; note those curves are each share-normalised, so they show timing, not the volume balance "
       "(which is net_dir).", False)])

H("5.1 Rail — weekday (MON+TWT, K=5, tied)", size=12)
sig_table("rail_weekday")
fig(FIG / "rail_weekday_k5_profiles_tail.png", width=5.6, cap="Fig. 3 · Rail weekday K=5 profiles (right column zooms the post-midnight tail)")
fig(FIG / "rail_weekday_k5_map.png", width=5.4, cap="Fig. 4 · Rail weekday K=5 — station clusters")

H("5.2 Rail — weekend (Fri+Sat, Night Tube, K=4, tied)", size=12)
para([("The cleanest typology. A small central nightlife ORIGIN core (C2, n=15: Tottenham Court Road, Leicester "
       "Square, Piccadilly Circus — the West End) carries the highest volume and the strongest 22–23h hump and "
       "is entry-dominant (people leaving after a night out). Two destination masses split by Night-Tube intensity "
       "(C0 overnight-active vs C1 limited), plus the major termini (C3).", False)])
sig_table("rail_weekend")
fig(FIG / "rail_weekend_k4_profiles_tail.png", width=5.6, cap="Fig. 5 · Rail weekend K=4 profiles (right column zooms the post-midnight tail)")
fig(FIG / "rail_weekend_k4_map.png", width=5.4, cap="Fig. 6 · Rail weekend K=4 — C2 (nightlife) concentrates in the Zone-1 core")

H("5.3 Bus — weekday (K=4, full)", size=12)
sig_table("bus_weekday")
fig(FIG / "bus_weekday_k4_profiles_tail.png", width=5.6, cap="Fig. 7 · Bus weekday K=4 profiles (boardings green / alightings red; right column zooms post-midnight)")
fig(FIG / "bus_weekday_k4_map.png", width=5.4, cap="Fig. 8 · Bus weekday K=4 — LSOA clusters")

H("5.4 Bus — weekend (Sat+Sun, K=4, full)", size=12)
sig_table("bus_weekend")
fig(FIG / "bus_weekend_k4_profiles_tail.png", width=5.6, cap="Fig. 9 · Bus weekend K=4 profiles (right column zooms post-midnight)")
fig(FIG / "bus_weekend_k4_map.png", width=5.4, cap="Fig. 10 · Bus weekend K=4 — LSOA clusters")

# ===================== 6 CROSS-CUTTING =====================
H("6. Cross-cutting findings")
bullet([("Origin vs destination is the primary night-time axis ", True), ("(consistent across rail and bus): outflow "
        "= activity/employment districts emptying; inflow = residential areas receiving returnees.", False)])
bullet([("A recurring small 'nightlife + airport' special type ", True), ("appears in both modes (rail weekend C2 / "
        "West End; bus C3 weekday and C0 weekend / Hatton Cross–Heathrow), always origin-dominant with the highest "
        "overnight share — the fingerprint of the night-time economy and airport shift travel.", False)])
bullet([("Bus surfaces a hospital / night-shift destination cluster ", True), ("(weekday C1: King's College "
        "Hospital and other hospitals, overnight-active, destination-dominant) — a direct hook to RQ2 (night "
        "workers / LNWC), more so than rail.", False)])
bullet([("Supply vs demand caveat (important): ", True), ("the rail-weekend overnight split partly reflects which "
        "stations lie on a Night-Tube LINE (overnight service runs on only five lines); off-network stations have "
        "structurally zero overnight activity. The 'overnight present/absent' cut therefore measures network "
        "coverage (supply) as much as demand — central to the project's supply-constrained framing and RQ3.", False)])

# ===================== 7 LIMITATIONS =====================
H("7. Scope and limitations")
for txt in [
    "Rail covers the Underground only (MON+TWT vs Fri+Sat); Sunday is excluded for now.",
    "Bus is LSOA-level (MAUP): a smooth continuum with weak silhouettes; clusters are intensity/overnight gradients rather than sharp functional types.",
    "All features are scale-free shares; absolute volume is retained separately but not yet brought back as an equity layer (cluster × volume tier).",
    "Counts are typical-week averages (NUMBAT/BUSTO); they measure observed, supply-constrained USE, not latent demand or residence.",
    "Cluster names are tentative (the code uses cluster ids only) and to be finalised at the meeting.",
]:
    bullet(txt)

# ===================== 8 OPEN =====================
H("8. For discussion")
for txt in [
    "Whether to bring Sunday back as a third rail regime, and whether to split bus weekend into Saturday vs Sunday.",
    "K confirmation per dataset (BIC=2 vs the interpretable K used).",
    "Overlaying Night-Tube line membership to separate the supply (coverage) and demand components of the overnight split.",
    "Carrying the typologies into RQ2 (LNWC / socio-economic overlay) — the hospital/night-shift bus cluster is the natural starting point.",
]:
    bullet(txt)

doc.save(str(DEST))
print("saved:", DEST)
